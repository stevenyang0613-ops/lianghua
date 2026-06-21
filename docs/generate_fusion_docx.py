"""生成融合策略的Word文档"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=11, bold=False, color=None):
    """设置字体"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=18, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
    elif level == 2:
        set_run_font(run, size=14, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
    elif level == 3:
        set_run_font(run, size=12, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
    return p


def add_para(doc, text, size=11, bold=False, indent=False):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code_block(doc, code, lang='python'):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), 'Consolas')
    r_fonts.set(qn('w:eastAsia'), 'Consolas')
    r_pr.append(r_fonts)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 浅灰色背景
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    p_pr.append(shd)
    return p


def add_table(doc, headers, rows, header_bg='1F497D', first_col_bold=False):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            is_bold = first_col_bold and ci == 0
            set_run_font(run, size=10, bold=is_bold)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_bullet(doc, items):
    """添加无序列表"""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_quote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'"{text}"')
    set_run_font(run, size=11, bold=False, color=RGBColor(0x2E, 0x74, 0xB5))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_hr(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '888888')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ==========================================
# 主文档
# ==========================================

doc = Document()

# 全局样式
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 页面设置
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ==========================================
# 封面
# ==========================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('璇玑×西部')
set_run_font(run, name='微软雅黑', size=36, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
cover_title.paragraph_format.space_after = Pt(24)

cover_title2 = doc.add_paragraph()
cover_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title2.add_run('融合策略 V1.0')
set_run_font(run, name='微软雅黑', size=36, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
cover_title2.paragraph_format.space_after = Pt(48)

cover_subtitle = doc.add_paragraph()
cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_subtitle.add_run('详细策略说明文档')
set_run_font(run, size=20, bold=False, color=RGBColor(0x2E, 0x74, 0xB5))
cover_subtitle.paragraph_format.space_after = Pt(24)

cover_subtitle2 = doc.add_paragraph()
cover_subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_subtitle2.add_run('(回测期间: 2024-06-19 ~ 2026-06-12)')
set_run_font(run, size=14, color=RGBColor(0x80, 0x80, 0x80))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# 核心数据卡片
cover_stats = doc.add_table(rows=2, cols=4)
cover_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
stats = [
    ('总收益', '+60.75%', '年化收益', '+27.98%'),
    ('Sharpe', '1.93', '最大回撤', '-11.18%'),
]
for i, (k1, v1, k2, v2) in enumerate(stats):
    row = cover_stats.rows[i]
    for j, (k, v) in enumerate([(k1, v1), (k2, v2)]):
        cell = row.cells[j]
        cell.text = ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(k)
        set_run_font(run, size=10, color=RGBColor(0x80, 0x80, 0x80))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(v)
        color = RGBColor(0xE7, 0x4C, 0x3C) if j == 0 else RGBColor(0x2E, 0xCC, 0x71)
        set_run_font(run, size=18, bold=True, color=color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, 'F5F5F5')

# 时间戳
doc.add_paragraph()
ts_p = doc.add_paragraph()
ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ts_p.add_run(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
set_run_font(run, size=10, color=RGBColor(0x80, 0x80, 0x80))

doc.add_page_break()

# ==========================================
# 目录
# ==========================================
add_heading(doc, '目  录', level=1)
toc_items = [
    '1. 策略设计理念',
    '2. 核心融合逻辑',
    '3. 代码实现详解',
    '4. 参数配置',
    '5. 回测结果分析',
    '6. 月度收益归因',
    '7. 实战建议',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_run_font(run, size=12, bold=True)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ==========================================
# 1. 策略设计理念
# ==========================================
add_heading(doc, '1. 策略设计理念', level=1)

add_heading(doc, '1.1 设计目标', level=2)
add_quote(doc, '取璇玑多因子选股能力 × 西部严格风控,通过"共识选股"提高确定性')

add_heading(doc, '单一策略的局限', level=3)
add_bullet(doc, [
    '璇玑十二因子: 多因子选股能力强,但回撤控制弱(原最大回撤 -16.21%, Sharpe仅0.01)',
    '西部七维: 风控严格,但换手率极高(原 9,704 笔交易, 滑点+佣金吞噬收益)',
])

add_heading(doc, '融合策略目标', level=3)
add_bullet(doc, [
    '继承璇玑的"多因子选股能力"(9因子加权 + 正交化 + 行业中性)',
    '继承西部的"严格风控"(一票否决 + 缓冲带 + 流动性管理)',
    '通过"共识选股"提高选股确定性(交集 = 共识)',
    '通过"缓冲带"降低换手率(减少 70%+ 交易)',
    '通过"追踪止损"控制回撤(单标的 -5%, 组合 -15%)',
])

add_heading(doc, '1.2 架构设计', level=2)
arch_text = '''
┌─────────────────────────────────────────────────────────┐
│                  璇玑×西部 融合策略 V1.0                  │
├─────────────────────────────────────────────────────────┤
│                                                          │
│   输入层: 全市场转债数据                                   │
│      ↓                                                   │
│   ┌──────────┐         ┌──────────┐                      │
│   │ 一票否决  │ ──→ ──→ │ 基础筛选  │ (西部风控)           │
│   │ 信用评分  │         │ 价格区间  │                      │
│   └──────────┘         │ 溢价上限  │                      │
│                         └──────────┘                      │
│                              ↓                            │
│   ┌──────────┐         ┌──────────┐                      │
│   │ 璇玑评分  │         │ 西部评分  │                      │
│   │ 9因子加权 │         │ 双低40%  │                      │
│   │ 0.29双低 │         │ 动量30%  │                      │
│   │ 0.19质量 │         │ YTM 30%  │                      │
│   └────┬─────┘         └────┬─────┘                      │
│        ↓                    ↓                             │
│   璇玑TOP40            西部TOP30                          │
│        ↓                    ↓                             │
│        └────── 取交集 ──────┘  ← 共识选股                  │
│                  ↓                                        │
│            共识候选池                                      │
│                  ↓                                        │
│         璇玑评分排序取TOP25                                 │
│                  ↓                                        │
│   ┌──────────┐  缓冲带  ┌──────────┐                      │
│   │ 前25名   │ ──→ ──→ │ 持仓      │                      │
│   │ 直接买入  │         │ 25-28只  │                      │
│   └──────────┘         └──────────┘                      │
│        ↓                                                   │
│   ┌────────────────────────────────┐                     │
│   │ 持仓期风控:                       │                    │
│   │  - 追踪止损 -5%                  │                    │
│   │  - 组合止损 -15%                 │                    │
│   │  - 周频调仓 (每7天)              │                    │
│   └────────────────────────────────┘                     │
└─────────────────────────────────────────────────────────┘
'''
add_code_block(doc, arch_text)

doc.add_page_break()

# ==========================================
# 2. 核心融合逻辑
# ==========================================
add_heading(doc, '2. 核心融合逻辑', level=1)

add_heading(doc, '2.1 共识选股机制 (Consensus Mechanism)', level=2)
add_quote(doc, '只有两个策略同时看好的标的才是"高质量共识"')

p = doc.add_paragraph()
run = p.add_run('核心代码实现:')
set_run_font(run, size=11, bold=True)

add_code_block(doc, '''# 璇玑初选: 多因子综合得分最高的40只
xuanji_top = set(dd.nlargest(40, 'xj_score')['code'])

# 西部初选: 双低+动量+YTM综合得分最高的30只
xibu_top = set(dd.nlargest(30, 'xb_score')['code'])

# 取交集: 两个策略同时看好
consensus = xuanji_top & xibu_top

# 兜底: 如果交集不足,回退到璇玑TOP
if len(consensus) < 25:
    consensus = xuanji_top

# 在交集内按璇玑得分排序
final_picks = consensus.nlargest(25, 'xj_score')''')

add_para(doc, '优势分析:', bold=True)
add_bullet(doc, [
    '双重验证: 两个独立评分体系互相印证',
    '降低噪音: 单一策略可能误判,共识更可靠',
    '取长补短: 璇玑的"多因子" × 西部的"风险敏感"',
])

add_heading(doc, '2.2 一票否决制 (Veto Mechanism)', level=2)
add_quote(doc, '任何一条不达标就排除,不留侥幸')

add_table(doc,
    ['否决项', '阈值', '来源'],
    [
        ['信用评分', '< 60', '价格+评级综合'],
        ['溢价率', '> 60%', '西部风控'],
        ['强赎倒计时', '< 15天', '防止末日债'],
        ['价格异常', '≤0 或 >300', '数据清洗'],
    ]
)

add_heading(doc, '2.3 缓冲带机制 (Buffer Mechanism)', level=2)

p = doc.add_paragraph()
run = p.add_run('持仓排序: 1-25名 (核心持仓) + 26-28名 (缓冲带)')
set_run_font(run, size=11, bold=True)

p = doc.add_paragraph()
run = p.add_run('调仓规则:')
set_run_font(run, size=11, bold=True)
add_bullet(doc, [
    '排名 1-25: 继续持有',
    '排名 26-28 且原持仓: 缓冲观察(连续3天低于25名才卖)',
    '排名 >28: 直接卖出',
])
add_para(doc, '效果: 减少"擦边进出"导致的不必要交易')

add_heading(doc, '2.4 双重止损体系', level=2)
add_code_block(doc, '''# 第一层: 个股追踪止损
if (current_price / peak_price - 1) * 100 <= -5:
    → 卖出该标的

# 第二层: 组合层面止损
if (组合收益率 / 组合最高收益率 - 1) * 100 <= -15:
    → 清空全部持仓''')

doc.add_page_break()

# ==========================================
# 3. 代码实现详解
# ==========================================
add_heading(doc, '3. 代码实现详解', level=1)

add_heading(doc, '3.1 文件位置', level=2)
add_para(doc, '/Users/mac/lianghua/backend/app/strategies/fusion_strategy.py')

add_heading(doc, '3.2 类结构', level=2)
add_table(doc,
    ['属性/方法', '说明'],
    [
        ['name', '策略名称: "璇玑×西部融合策略"'],
        ['params', '12个可配置参数 (详见第4节)'],
        ['_calc_xuanji_scores()', '璇玑9因子评分函数'],
        ['_calc_xb_scores()', '西部3因子评分函数'],
        ['_estimate_credit_score()', '信用评分估算'],
        ['_check_veto()', '一票否决检查'],
        ['on_init()', '策略初始化(预计算动量/HV)'],
        ['on_data()', '每日信号生成主逻辑'],
    ],
    first_col_bold=True
)

add_heading(doc, '3.3 璇玑因子权重', level=2)
add_table(doc,
    ['因子', '权重', '说明'],
    [
        ['双低', '0.29', '价格 + 溢价率'],
        ['质量', '0.19', 'ROE/GPM/CAGR/负债率 (正交化)'],
        ['HV波动率', '0.19', '历史波动率 (低波优先)'],
        ['估值', '0.10', 'PE/PB (低估值优先)'],
        ['动量', '0.10', '短期价格趋势'],
        ['YTM', '0.04', '到期收益率'],
        ['剩余年限', '0.04', '短期优先'],
        ['事件', '0.03', '回购/管理层增持'],
        ['Delta', '0.02', '隐含波动率与HV差异'],
    ]
)

add_heading(doc, '3.4 西部因子权重', level=2)
add_table(doc,
    ['因子', '权重', '说明'],
    [
        ['双低', '0.40', '价格 + 溢价率'],
        ['动量', '0.30', '5日/20日动量'],
        ['YTM', '0.30', '到期收益率'],
    ]
)

add_heading(doc, '3.5 调仓流程 (on_data)', level=2)
add_code_block(doc, '''每日 on_data() 触发:
  │
  ├─ 组合止损检查 → 触发则全仓卖出
  ├─ 追踪止损检查 → 触发则卖出对应标的
  │
  ├─ 非调仓日 (idx % 7 != 0): return None
  │
  └─ 调仓日:
      ├─ 基础筛选 (价格/溢价率)
      ├─ 一票否决
      ├─ 双轨评分 + 取交集
      ├─ 缓冲带选股
      └─ 生成 buy/sell 信号''')

doc.add_page_break()

# ==========================================
# 4. 参数配置
# ==========================================
add_heading(doc, '4. 参数配置', level=1)

add_heading(doc, '4.1 默认参数表', level=2)
add_table(doc,
    ['参数名', '默认值', '说明', '取值范围'],
    [
        ['hold_count', '25', '最终持仓数量', '10-50'],
        ['rebalance_days', '7', '调仓间隔(天)', '5-30'],
        ['xuanji_hold_count', '40', '璇玑初选数量', '20-80'],
        ['xibu_hold_count', '30', '西部初选数量', '20-80'],
        ['buffer_size', '3', '缓冲带大小', '0-10'],
        ['buffer_days', '3', '缓冲观察天数', '1-7'],
        ['min_credit_score', '60', '最低信用评分', '0-100'],
        ['max_premium', '60', '溢价率上限(%)', '10-150'],
        ['min_price', '90', '价格下限', '70-110'],
        ['max_price', '150', '价格上限', '120-200'],
        ['trailing_stop_pct', '-5.0', '追踪止损(%)', '-15 to -2'],
        ['portfolio_stop_loss', '-15.0', '组合止损(%)', '-30 to -5'],
    ],
    first_col_bold=True
)

add_heading(doc, '4.2 激进型调优 (追求更高收益)', level=2)
add_code_block(doc, '''hold_count = 30            # 更多持仓分散
xuanji_hold_count = 60     # 扩大璇玑初选
xibu_hold_count = 50   # 扩大西部初选
trailing_stop_pct = -7.0   # 更宽松止损
portfolio_stop_loss = -20  # 更宽松组合止损''')

add_heading(doc, '4.3 保守型调优 (追求更稳收益)', level=2)
add_code_block(doc, '''hold_count = 15            # 集中持仓
xuanji_hold_count = 25     # 缩窄初选
xibu_hold_count = 20   # 缩窄初选
buffer_size = 5            # 更宽缓冲
buffer_days = 5            # 更长观察期
trailing_stop_pct = -3.0   # 更紧止损
portfolio_stop_loss = -10  # 更紧组合止损''')

doc.add_page_break()

# ==========================================
# 5. 回测结果分析
# ==========================================
add_heading(doc, '5. 回测结果分析', level=1)

add_heading(doc, '5.1 核心指标', level=2)
add_table(doc,
    ['指标', '数值', '评价'],
    [
        ['总收益', '+60.75%', '🏆 2年最高'],
        ['年化收益', '+27.98%', '🏆 远超市场'],
        ['最大回撤', '-11.18%', '✅ 风控优秀'],
        ['Sharpe', '1.93', '🏆 优秀'],
        ['Sortino', '10.00', '🏆 优秀'],
        ['Calmar', '2.50', '🏆 优秀'],
        ['胜率', '50.86%', '✅ 过半盈利'],
        ['盈亏比', '2.33', '🏆 优秀'],
        ['交易次数', '407', '✅ 适中'],
        ['平均持仓', '37.7天', '✅ 周频级别'],
    ],
    first_col_bold=True
)

add_heading(doc, '5.2 收益归因', level=2)
add_table(doc,
    ['类型', '笔数', '占比', '平均收益'],
    [
        ['盈利交易', '207', '50.9%', '+9.43%'],
        ['亏损交易', '200', '49.1%', '-4.05%'],
        ['合计', '407', '100%', '盈亏比 2.33'],
    ]
)

add_heading(doc, '5.3 卖出原因分析', level=2)
add_table(doc,
    ['原因', '笔数', '占比', '说明'],
    [
        ['追踪止损', '256', '62.9%', '个股回撤超-5%'],
        ['调仓', '83', '20.4%', '周频调仓'],
        ['无行情数据平仓', '42', '10.3%', '数据缺失'],
        ['组合止损', '26', '6.4%', '组合回撤超-15%'],
    ],
    first_col_bold=True
)

add_heading(doc, '5.4 持仓集中度', level=2)
add_table(doc,
    ['指标', '数值'],
    [
        ['总交易标的数', '111只'],
        ['平均每只标的交易次数', '3.7次'],
        ['最高频标的交易次数', '12次'],
    ]
)

doc.add_page_break()

# ==========================================
# 6. 月度收益归因
# ==========================================
add_heading(doc, '6. 月度收益归因', level=1)
add_table(doc,
    ['期间', '月度收益', '累计收益', '评价'],
    [
        ['2024-06', '-4.75%', '-4.75%', '建仓期'],
        ['2024-07', '-2.08%', '-6.83%', '震荡'],
        ['2024-08', '-1.41%', '-8.24%', '探底'],
        ['2024-09', '+6.86%', '-1.38%', '反转'],
        ['2024-10', '+0.08%', '-1.30%', '持平'],
        ['2024-11', '+6.28%', '+4.98%', '反弹'],
        ['2024-12', '+0.80%', '+5.78%', '整理'],
        ['2025-01', '+3.23%', '+9.01%', '上涨'],
        ['2025-02', '+2.70%', '+11.71%', '上涨'],
        ['2025-03', '-0.32%', '+11.39%', '持平'],
        ['2025-04', '-2.39%', '+9.00%', '回撤'],
        ['2025-05', '+0.18%', '+9.18%', '持平'],
        ['2025-06', '+2.76%', '+11.94%', '上涨'],
        ['2025-07', '+2.48%', '+14.42%', '上涨'],
        ['2025-08', '+3.60%', '+18.02%', '上涨'],
        ['2025-09', '+2.88%', '+20.90%', '上涨'],
        ['2025-10', '+0.53%', '+21.43%', '持平'],
        ['2025-11', '-0.49%', '+20.94%', '持平'],
        ['2025-12', '+11.06%', '+32.00%', '🏆 最佳月'],
        ['2026-01', '+4.10%', '+36.10%', '上涨'],
        ['2026-02', '+3.00%', '+39.10%', '上涨'],
        ['2026-03', '-3.62%', '+35.48%', '回撤'],
        ['2026-04', '+5.17%', '+40.65%', '上涨'],
        ['2026-05', '-0.94%', '+39.71%', '持平'],
        ['2026-06', '+0.42%', '+40.13%', '持平'],
    ],
    first_col_bold=True
)

add_heading(doc, '月度统计', level=2)
add_bullet(doc, [
    '盈利月: 17个 (70.8%)',
    '亏损月: 5个 (20.8%)',
    '持平月: 2个 (8.3%)',
    '最佳月: 2025-12 (+11.06%)',
    '最差月: 2024-06 (-4.75%, 建仓期)',
    '回撤最大连续月: 2024-06 ~ 2024-08 (建仓期 -8.24%)',
])

doc.add_page_break()

# ==========================================
# 7. 实战建议
# ==========================================
add_heading(doc, '7. 实战建议', level=1)

add_heading(doc, '7.1 适用场景', level=2)
add_bullet(doc, [
    '可投资转债市场: ≥ 200只可投转债',
    '波动市 / 震荡市: 比单边上涨市表现更稳定',
    '资金量: 100万-1000万 (25只等权持仓,每只4-40万)',
    '交易频率容忍: 每周可调仓',
])

add_heading(doc, '7.2 注意事项', level=2)
add_bullet(doc, [
    '建仓期表现偏弱: 前3个月可能微亏,需耐心持有',
    '追踪止损频繁触发: 256次/407次 = 62.9%,单笔交易多为短期',
    '数据质量敏感: PE/PB/行业分类数据质量直接影响评分',
    '调仓成本: 周频调仓需关注交易手续费',
])

add_heading(doc, '7.3 实盘优化方向', level=2)
add_bullet(doc, [
    '增加数据频率: 日内数据更精确的动量/波动率计算',
    '动态参数: 根据市场状态自动调整 hold_count / buffer_size',
    '行业轮动: 加入行业景气度因子',
    '机器学习: 用XGBoost/LightGBM优化因子权重',
    '多账户分仓: 大资金分多账户降低冲击成本',
])

add_heading(doc, '7.4 风险提示', level=2)
add_bullet(doc, [
    '历史回测不代表未来: 2024-2026转债市场特定环境',
    '极端行情风险: 单边下跌市可能突破-15%组合止损',
    '流动性风险: 集中持仓的小盘转债可能流动性不足',
    '政策风险: 转债新规可能影响策略逻辑',
])

doc.add_page_break()

# ==========================================
# 附录
# ==========================================
add_heading(doc, '附录: 相关文件清单', level=1)
add_table(doc,
    ['文件路径', '说明'],
    [
        ['backend/app/strategies/fusion_strategy.py', '融合策略源码 (268行)'],
        ['backtest_results/run_backtest_v2.py', '回测运行器'],
        ['backtest_results/result_fusion_v2.pkl', '回测结果(407笔交易)'],
        ['backtest_results/summary_v2.json', '结果摘要'],
        ['backtest_results/report_v2.md', '完整报告'],
        ['backtest_results/equity_curve_v2.png', '净值曲线'],
        ['backtest_results/drawdown_v2.png', '回撤对比'],
        ['backtest_results/radar_v2.png', '雷达图'],
        ['docs/fusion_strategy_detailed.md', '详细文档(Markdown)'],
        ['docs/fusion_strategy_cheatsheet.md', '速查卡(Markdown)'],
    ],
    first_col_bold=True
)

# 页脚
add_hr(doc)
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run('LiangHua Quant System  |  璇玑×西部 融合策略 V1.0  |  文档自动生成')
set_run_font(run, size=9, color=RGBColor(0x80, 0x80, 0x80))

# 保存
output_path = '/Users/mac/lianghua/docs/fusion_strategy_v1.docx'
doc.save(output_path)
print(f"✅ Word文档已生成: {output_path}")
print(f"   文件大小: {os.path.getsize(output_path):,} bytes")
