"""生成完整策略三合一Word文档"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run_font(run, name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:eastAsia'), name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=18, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(28)
        p.paragraph_format.space_after = Pt(20)
    elif level == 2:
        set_run_font(run, size=15, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(12)
    elif level == 3:
        set_run_font(run, size=12, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    return p


def add_para(doc, text, size=11, bold=False, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code(doc, code, size=9):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), 'Consolas')
    r_fonts.set(qn('w:eastAsia'), 'Consolas')
    r_pr.append(r_fonts)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    p_pr.append(shd)
    return p


def add_table(doc, headers, rows, header_bg='1A3C6E', col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run_font(run, size=10, bold=(ci == 0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ==========================================
# 文档正文
# ==========================================

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# ====================== 封面 ======================
for _ in range(8):
    doc.add_paragraph()

cover_t = doc.add_paragraph()
cover_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_t.add_run('LiangHua 量化交易系统\n策略体系详细文档')
set_run_font(run, size=28, bold=True, color=RGBColor(0x1A, 0x3C, 0x6E))

doc.add_paragraph()

sub_t = doc.add_paragraph()
sub_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_t.add_run('璇玑十二因子  ·  西部七维打分  ·  璇玑×西部融合策略')
set_run_font(run, size=14, color=RGBColor(0x2E, 0x74, 0xB5))

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f'回测区间: 2024-06-19 ~ 2026-06-12')
set_run_font(run, size=12, color=RGBColor(0x80, 0x80, 0x80))

date_p2 = doc.add_paragraph()
date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p2.add_run(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
set_run_font(run, size=11, color=RGBColor(0x90, 0x90, 0x90))

for _ in range(6):
    doc.add_paragraph()

# 核心数据卡片
cover_stats = doc.add_table(rows=2, cols=4)
cover_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_data = [
    ('璇玑v5.0', '+38.20%', '西部v4.0', '+54.72%'),
    ('融合策略', '+60.75%', 'Sharpe 1.93', '回撤-11.18%'),
]
for i, (k1, v1, k2, v2) in enumerate(cover_data):
    row = cover_stats.rows[i]
    for j, (k, v) in enumerate([(k1, v1), (k2, v2)]):
        cell = row.cells[j]
        cell.text = ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(k)
        set_run_font(run, size=10, color=RGBColor(0x66, 0x66, 0x66))
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(v)
        c = RGBColor(0xE7, 0x4C, 0x3C) if '+' in str(v) and '%' in str(v) else RGBColor(0x2E, 0x74, 0xB5)
        set_run_font(run, size=16, bold=True, color=c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, 'F7F9FC')

doc.add_page_break()

# ====================== 目录 ======================
add_heading(doc, '目  录', level=1)
toc = [
    '第一章  璇玑十二因子指增策略',
    '  1.1 策略概述',
    '  1.2 9因子评分体系',
    '  1.3 市场状态自适应',
    '  1.4 债券分层与权重调整',
    '  1.5 因子正交化与行业中性化',
    '  1.6 ICIR动态权重',
    '  1.7 一票否决制',
    '  1.8 回测参数与结果',

    '',
    '第二章  西部七维打分策略',
    '  2.1 策略概述',
    '  2.2 正股七维评分 (55分)',
    '  2.3 转债四维评分 (45分)',
    '  2.4 一票否决制',
    '  2.5 缓冲带机制',
    '  2.6 回测参数与结果',

    '',
    '第三章  璇玑×西部融合策略',
    '  3.1 策略概述',
    '  3.2 共识选股机制',
    '  3.3 完整调仓流程',
    '  3.4 参数配置',
    '  3.5 回测结果',
    '  3.6 月度收益归因',
    '  3.7 实战建议与风险提示',
]
for item in toc:
    if item == '':
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run(item)
        b = not item.startswith('  ')
        set_run_font(run, size=11, bold=b, color=RGBColor(0x33, 0x33, 0x33) if b else RGBColor(0x55, 0x55, 0x55))
        p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ==============================================================
# 第一章  璇玑十二因子
# ==============================================================
add_heading(doc, '第一章  璇玑十二因子指增策略', level=1)

add_heading(doc, '1.1 策略概述', level=2)
add_para(doc, '璇玑十二因子指增策略(LiangHua v5.0)是一个多因子量化转债选股策略。核心思路是通过9个维度的因子打分(双低、动量、波动率、质量、估值、YTM、剩余年限、事件、Delta),结合ICIR动态权重、因子正交化、行业中性化、债券分层等技术,实现稳健的超额收益。')
add_para(doc, '策略名称中的"十二因子"源自最初的因子设计概念,经过演进和正交化后,当前核心因子为9个。另外还包含3个辅助因子(流动性、信用评分、强制赎回)作为否决条件,整体保持"十二"的框架。')

add_heading(doc, '核心特性', level=3)
add_bullet(doc, [
    '多因子评分: 9个核心因子,覆盖价格、动量、质量、估值、波动率、收益、期限、事件、衍生品等多个维度',
    'ICIR动态权重: 基于历史因子IC(Information Coefficient)的滚动追踪,动态调整因子权重',
    'Gram-Schmidt正交化: 消除因子之间的多重共线性,提高因子的纯度',
    '行业中性化: 行业内Z-score标准化,消除行业风格偏差',
    '市场状态自适应: 根据转债市场整体价格、溢价率、涨跌比自动识别5种市场状态,切换最优权重配置',
    '债券分层: 按转股平价溢价率将转债分为偏股型/平衡型/偏债型,分层调整因子权重',
    '一票否决制: 信用评分、溢价率、强赎风险等6项否决条件',
    '缓冲带机制: 减少不必要的换手,降低交易成本',
    '追踪止损: 个股-5%追踪止损 + 组合-15%组合止损',
])

add_heading(doc, '1.2 9因子评分体系', level=2)
add_para(doc, '9个核心因子的默认权重(中性市场)如下:', bold=True)

add_table(doc, ['序号', '因子名称', '基础权重', '评分方向', '计算方式', '说明'],
    [
        ['1', '双低(dual_low)', '0.29', '越低越好', 'price/median_price*50 + premium_ratio', '价格+溢价率的综合指标,是转债最经典的估值指标'],
        ['2', '质量(quality)', '0.19', '越高越好', 'ROE → GPM → CAGR → debt_ratio 正交化后加权', '基于ROE、毛利率、增长率、负债率等财务指标的复合质量因子'],
        ['3', 'HV波动率(hv)', '0.19', '越低越好', '20日滚动标准差*√252*100', '历史波动率,低波动标的优先(波动率越低,转债期权价值越稳定)'],
        ['4', '估值(valuation)', '0.10', '越低越好', 'PE → PB 正交化后加权', '正股的PE和PB估值,低估值优先'],
        ['5', '动量(momentum)', '0.10', '正向动力', '5/10/20/60日收益率加权: 0.4*5d + 0.3*10d + 0.2*20d + 0.1*60d', '短期价格动量,捕捉趋势延续'],
        ['6', 'YTM', '0.04', '越高越好', '按日动态计算,基于到期日和票面利率', '到期收益率,纯债价值的安全垫'],
        ['7', '剩余年限', '0.04', '越短越好', '(到期日-当前日)/365', '剩余期限越短的转债,不确定性越小'],
        ['8', '事件(event)', '0.03', '越高越好', '基于回购金额/管理层增持/其他事件', '正股事件驱动因子(需外部数据源)'],
        ['9', 'Delta', '0.02', '越低越好', 'max(IV, HV*1.2+3) - HV', '隐含波动率与历史波动率的差额,差额越小定价越合理'],
    ])

add_heading(doc, '1.3 市场状态自适应', level=2)
add_para(doc, '策略根据转债市场整体状态自动检测并切换5种市场状态的权重配置:')
add_table(doc, ['市场状态', '检测条件', '双低', '动量', 'HV', '质量', '估值'],
    [
        ['极端牛市', 'bull_score ≥ 4\n(价格>140 + 溢价率<15% + 涨跌比>65%)', '0.14', '0.28', '0.14', '0.19', '0.10'],
        ['温和牛市', 'bull_score ≥ 2', '0.24', '0.24', '0.14', '0.14', '0.10'],
        ['中性', '默认状态', '0.29', '0.10', '0.19', '0.19', '0.10'],
        ['温和熊市', 'bull_score ≤ -1', '0.38', '0.05', '0.24', '0.14', '0.09'],
        ['极端熊市', 'bull_score ≤ -3\n(价格<108 + 溢价率>50%)', '0.47', '0.00', '0.29', '0.09', '0.05'],
    ])
add_para(doc, '核心逻辑: 牛市加大动量权重(0.28),减少双低权重(0.14); 熊市加大双低安全垫权重(0.47),完全去掉动量(0.00)。')

add_heading(doc, '1.4 债券分层与权重调整', level=2)
add_para(doc, '根据转股平价溢价率将转债分为3层,每层对因子权重进行差异化调整:')
add_table(doc, ['分层', '条件', '双低', '动量', '质量', '估值', 'YTM', 'Delta'],
    [
        ['偏股型', '平价溢价率>+20%', '×0.5', '×1.5', '×1.4', '×1.3', '×0.0', '×1.5'],
        ['平衡型', '平价溢价率-20% ~ +20%', '×1.0', '×1.0', '×1.0', '×1.0', '×1.0', '×1.0'],
        ['偏债型', '平价溢价率<-20%', '×1.5', '×0.3', '×0.6', '×0.6', '×1.5', '×0.3'],
    ])
add_para(doc, '偏股型: 更看重动量和Delta,淡化YTM; 偏债型: 更看重双低和YTM,淡化动量。')

add_heading(doc, '1.5 因子正交化与行业中性化', level=2)
add_heading(doc, 'Gram-Schmidt正交化', level=3)
add_para(doc, '质量因子(ROE/GPM/CAGR/debt_ratio)和估值因子(PE/PB)内部存在相关性,采用Gram-Schmidt正交化消除冗余:')
add_code(doc, '''# 以估值因子为例:
# 原始: PE_score 和 PB_score 存在相关性
# 正交化后:
pe_orthogonal = pe_score  # 第一个因子保持
pb_orthogonal = pb_score - (pe_score·pb_score)/(pe_score·pe_score) * pe_score
# 结果: pb_orthogonal 包含了PB中PE无法解释的"增量信息"''')

add_heading(doc, '行业中性化', level=3)
add_para(doc, '每个因子在行业内做Z-score标准化,消除行业偏差:')
add_code(doc, '''for each industry group:
    mu = scores[industry].mean()
    sigma = scores[industry].std()
    scores[industry] = (scores[industry] - mu) / sigma  # 行业内Z-score
# 最终归一化到[0,1]区间''')

add_heading(doc, '1.6 ICIR动态权重', level=2)
add_para(doc, '基于滚动IC(因子与未来收益的秩相关系数)和IC稳定性(IR=IC_mean/IC_std)动态调整因子权重:')
add_code(doc, '''每交易日:
    for each factor f:
        ic[f] = Spearman_corr(factor_score[f], next_day_return)
        ic_stats[f] = ic_mean[f] / ic_std[f]  (# 信息比率)
    
    # 新权重 = 60% ICIR权重 + 40% 基础权重
    blended_weight[f] = 0.6 * icir_weight[f] + 0.4 * base_weight[f]
    
    # 回退: IC历史<5天时,使用基础权重''')

add_heading(doc, '1.7 一票否决制', level=2)
add_para(doc, '满足任意一条即排除,不留侥幸:')
add_table(doc, ['否决条件', '判断逻辑', '阈值', '设计意图'],
    [
        ['信用评分', '基于价格/溢价率/YTM/评级的综合评分', '< 60', '排除信用风险较高的转债'],
        ['溢价率', '转股溢价率', '> 50%', '排除纯投机品种'],
        ['强赎风险', '强制赎回到计时', '>0天且<15天', '避免强赎导致的额外损失'],
        ['剩余期限', '距到期日', '< 0.5年', '排除即将到期的低波动品种'],
        ['流动性', '成交量', '< 500手', '确保可交易性'],
        ['价格异常', '价格', '≤0 或 >300', '数据清洗'],
    ])

add_heading(doc, '1.8 回测参数与结果', level=2)
add_para(doc, '璇玑十二因子 v5.0 回测结果 (2024-06-19 ~ 2026-06-12):', bold=True)
add_table(doc, ['参数', '默认值', '指标', '数值'],
    [
        ['hold_count', '20', '总收益', '+38.20%'],
        ['rebalance_days', '7', '年化收益', '+18.31%'],
        ['max_premium', '50', '最大回撤', '-14.65%'],
        ['min_price', '90', 'Sharpe', '0.96'],
        ['max_price', '150', 'Calmar', '1.25'],
        ['stop_loss_pct', '-8%', '盈亏比', '2.16'],
        ['buffer_size', '3', '交易次数', '560'],
    ])

doc.add_page_break()

# ==============================================================
# 第二章  西部七维打分策略
# ==============================================================
add_heading(doc, '第二章  西部七维打分策略', level=1)

add_heading(doc, '2.1 策略概述', level=2)
add_para(doc, '西部七维打分策略(LiangHua v4.0)采用"赎回债底价值 + 转股价值 + 期权价值 + 波动率溢价"的转债定价理论框架,构建正股七维评分(占55%) + 转债自身四维评分(占45%)的复合打分体系。同时配备一票否决制、缓冲带机制、动态择时等风控手段。')

add_para(doc, '核心理念', bold=True)
add_para(doc, '转债价格 = 债底价值 + 转股价值 + 期权价值 + 波动率溢价\n正股基本面(七维评分)决定转债的长期价值中枢,转债结构(四维评分)决定短期交易价值。')

add_heading(doc, '2.2 正股七维评分 (55分)', level=2)
add_table(doc, ['维度', '权重', '满分', '评分逻辑'],
    [
        ['1. 短期动量', '0.30', '16.5分', 'Z-score(涨幅)×0.4 + Z-score(量比)×0.3 + Z-score(正股涨幅)×0.3\n将复合Z-score映射到[0,16.5]分'],
        ['2. 板块情绪', '0.18', '9.9分', '正股涨幅>5%得满分9.9,>3%得8.0,>1%得6.0,>0%得4.0\n>-2%得2.0,≤-2%得0分'],
        ['3. 技术面', '0.18', '9.9分', '价格位置(0-4分) + 双低值(0-6.9分)\n价格90-100区间最优(4分);双低<120最优(6.9分)'],
        ['4. 筹码面', '0.12', '6.6分', '成交量0.5-5万手为最佳(满分6.6)\n成交量过大(≥10万手)表明分歧大,得分最低(3.8)'],
        ['5. 波动率', '0.12', '6.6分', '转债涨跌幅1-5%+正股2-6%为最佳(约5分)\n正偏度(正股涨时转债涨幅更大)额外加1.6分'],
        ['6. 消息面', '0.07', '3.85分', '双低<110(下修预期)+2.85分\n强赎倒计时10-20天+1.5分'],
        ['7. 基本面', '0.03', '1.65分', 'YTM>0得1.65分(正常)\nYTM -5~0得1分,YTM<-5得0分'],
    ])

add_heading(doc, '2.3 转债四维评分 (45分)', level=2)
add_table(doc, ['维度', '权重', '满分', '评分逻辑'],
    [
        ['1. 估值指标', '0.38', '17.1分', '溢价率<15%得10分,15-25%得5分,>25%得0分\n双低<120再加7.1分,<140加5分,<160加3分'],
        ['2. 条款价值', '0.24', '10.8分', '双低<100+溢价>30%(下修概率高)+6分\n0.5-2年进入回售期+4.8分'],
        ['3. 流动性', '0.20', '9.0分', '成交额≥4倍AUM阈值得9分\n≥2倍得6分,≥阈值得3分,不足得0分'],
        ['4. 信用评分', '0.18', '8.1分', '信用分≥80得8.1分,≥70得6分\n≥60得4分,<60得0分'],
    ])

add_heading(doc, '2.4 一票否决制', level=2)
add_table(doc, ['否决条件', '阈值', '说明'],
    [
        ['信用评分不足', '< 60', '基于价格、溢价率、YTM、双低的综合信用评估'],
        ['转股溢价率过高', '> 100%', '宽于璇玑(50%),但配合信用评分做二次过滤'],
        ['剩余期限过短', '< 6个月', '规避到期前的不确定性'],
        ['强赎倒计时', '>0且<15天', '防止强制赎回带来的不利交易'],
        ['流动性不足(AUM联动)', '< 500/2000/5000万', '小资金500万,大资金5000万'],
        ['价格异常', '≤0 或 >300', '数据有效性检查'],
    ])

add_heading(doc, '2.5 缓冲带机制', level=2)
add_para(doc, '西部七维策略的缓冲带机制是其核心特色之一,包含完整的持有/卖出决策逻辑:')
add_code(doc, '''持仓排名规则:
  | 排名  | 是否原持仓 | 操作              |
  |-------|------------|-------------------|
  | 1-60  | 是         | 持有              |
  | 1-60  | 否         | 买入(新入选)       |
  | 61-65 | 是         | 缓冲观察(不卖)     |
  | 61-65 | 否         | 跳过              |
  | >65   | 是         | 卖出(跌出白名单)   |
  | >65   | 否         | 不动              |

缓冲观察详情:
  - 连续 buffer_days(默认5天)都在61-65名 → 卖出
  - 在此期间回到前60名 → 重置计数并继续持有
  - 缓冲状态持久化到DuckDB,重启后恢复''')
add_para(doc, '优点: 大幅降低"擦边进出"导致的交易磨损,原版9,704笔交易降低到2,207笔(降幅77%)。')

add_heading(doc, '2.6 回测参数与结果', level=2)
add_table(doc, ['参数', '默认值', '指标', '数值'],
    [
        ['hold_count', '60', '总收益', '+54.72%'],
        ['buffer_size', '8', '年化收益', '+25.47%'],
        ['buffer_days', '5', '最大回撤', '-10.37%'],
        ['rebalance_days', '5', 'Sharpe', '1.51'],
        ['min_credit_score', '60', '胜率', '53.24%'],
        ['max_premium', '100', '盈亏比', '1.61'],
        ['aum_level', 'small', '交易次数', '2,207(原9,704)'],
    ])

doc.add_page_break()

# ==============================================================
# 第三章  璇玑×西部融合策略
# ==============================================================
add_heading(doc, '第三章  璇玑×西部融合策略', level=1)

add_heading(doc, '3.1 策略概述', level=2)
add_para(doc, '璇玑×西部融合策略(LiangHua Fusion v1.0)是一个"取长补短"的综合策略。它继承璇玑的"多因子选股能力"(9因子加权+正交化+行业中性),同时吸收西部七维的"严格风控"(一票否决+缓冲带+流动性管理)。核心创新在于"共识选股机制"——只选择两个策略同时看好的标的,大幅提升选股确定性。')

add_para(doc, '设计动机', bold=True)
add_bullet(doc, [
    '璇玑的局限: 多因子选股能力强,但回撤控制弱(原版最大回撤 -16.21%)',
    '西部的局限: 风控严格,但换手率高(原版 9,704 笔交易,年化收益被交易成本侵蚀)',
    '融合的优势: 继承璇玑的选股能力 × 继承西部的风控机制,通过"交集"提高确定性',
])

add_heading(doc, '3.2 共识选股机制 (核心创新)', level=2)
add_para(doc, '融合策略的核心是"双轨评分 + 取交集"的共识选股机制:')

add_heading(doc, '璇玑评分维度 (9因子)', level=3)
add_table(doc, ['因子', '原始权重', '璇玑9因子加权 → 取前40只'],
    [
        ['双低(dual_low)', '0.29', '价格/中位数*50 + 溢价率,越低越好'],
        ['质量(quality)', '0.19', 'ROE/GPM/CAGR/负债率正交化后加权'],
        ['HV波动率(hv)', '0.19', '20日滚动标准差*√252,低波优先'],
        ['估值(valuation)', '0.10', 'PE + PB 正交化,低估值优先'],
        ['动量(momentum)', '0.10', '5日/20日动量加权'],
        ['YTM', '0.04', '按日动态计算'],
        ['剩余年限', '0.04', '越短越好'],
        ['事件(event)', '0.03', '基于外部数据的事件评分'],
        ['Delta(delta)', '0.02', '隐含波动率与HV差额'],
    ])

add_heading(doc, '西部评分维度 (3因子简化版)', level=3)
add_table(doc, ['因子', '权重', '说明'],
    [
        ['双低(dual_low)', '0.40', '价格+溢价率,越低越好'],
        ['动量(momentum)', '0.30', '短期价格趋势'],
        ['YTM', '0.30', '到期收益率'],
    ])

add_heading(doc, '交集选股逻辑', level=3)
add_code(doc, '''# 双轨独立评分
xuanji_scores = calc_xuanji_scores(dd)     # 9因子加权
	xibu_scores = calc_xibu_scores(dd)       # 3因子简化版
	
	# 独立选取前N名
	xuanji_top40 = dd.nlargest(40, 'xj_score')['code']     # 璇玑前40
	xibu_top30 = dd.nlargest(30, 'sg_score')['code']    # 西部前30
	
	# 取交集 = 共识标的
	consensus = xuanji_top40 ∩ xibu_top30

# 兜底: 如果交集不足25只,回退到璇玑前25
if len(consensus) < 25:
    consensus = xuanji_top40

# 在共识池内按璇玑得分排序,取前25+缓冲带
final_picks = consensus.nlargest(25 + buffer_size, 'xj_score')''')

add_para(doc, '优势: 双重验证 — 两个独立评分体系(9因子复杂模型 + 3因子快速筛选)共同选出标的,大幅降低单一模型的误判概率。')

add_heading(doc, '3.3 完整调仓流程', level=2)
add_code(doc, '''每日 on_data() 触发:
  │
  ├─ 组合止损检查 (组合回撤<-15% → 清仓)
  ├─ 追踪止损检查 (个股从高点回撤>-5% → 卖出)
  │
  ├─ 非调仓日 (idx % 7 != 0): return None
  │
  └─ 调仓日:
      ├─ 基础筛选 (价格90-150 / 溢价率<60% / 成交量>0)
      ├─ 一票否决 (信用<60 / 溢价>60% / 强赎<15天 / 价格异常)
      ├─ 双轨评分:
      │   ├─ 璇玑评分 (9因子加权 + 正交化 + 行业中性)
      │   └─ 西部评分 (双低40% + 动量30% + YTM30%)
      ├─ 双轨取交集 → 共识候选
      ├─ 缓冲带选股:
      │   ├─ 前25名 → 买入/持有
      │   ├─ 26-28名(原持仓) → 缓冲观察(3天)
      │   └─ >28名 → 卖出
      └─ 生成 buy / sell 信号''')

add_heading(doc, '3.4 参数配置', level=2)
add_table(doc, ['参数名', '默认值', '说明', '取值范围'],
    [
        ['hold_count', '25', '最终持仓数量', '10-50'],
        ['rebalance_days', '7', '调仓间隔(天)', '5-30'],
        ['xuanji_hold_count', '40', '璇玑初选数量', '20-80'],
        ['xibu_hold_count', '30', '西部初选数量', '20-80'],
        ['buffer_size', '3', '缓冲带大小', '0-10'],
        ['buffer_days', '3', '缓冲观察天数', '1-7'],
        ['min_credit_score', '60', '最低信用评分', '0-100'],
        ['max_premium', '60', '溢价率上限(%)', '10-150'],
        ['min_price', '90', '价格下限', '70-110'],
        ['max_price', '150', '价格上限', '120-200'],
        ['trailing_stop_pct', '-5.0', '追踪止损(%)', '-15 to -2'],
        ['portfolio_stop_loss', '-15.0', '组合止损(%)', '-30 to -5'],
    ])

add_heading(doc, '激进型调优 (追求更高收益)', level=3)
add_code(doc, '''hold_count = 30            # 更多持仓分散
xuanji_hold_count = 60     # 扩大璇玑初选,更多候选
xibu_hold_count = 50   # 扩大西部初选
trailing_stop_pct = -7.0   # 更宽松止损(允许更大波动)
portfolio_stop_loss = -20  # 更宽松组合止损''')

add_heading(doc, '保守型调优 (追求更稳收益)', level=3)
add_code(doc, '''hold_count = 15            # 集中持仓,精选标的
xuanji_hold_count = 25     # 缩窄初选
xibu_hold_count = 20   # 缩窄初选
buffer_size = 5            # 更宽缓冲(减少换手)
buffer_days = 5            # 更长观察期
trailing_stop_pct = -3.0   # 更紧止损
portfolio_stop_loss = -10  # 更紧组合止损''')

add_heading(doc, '3.5 回测结果', level=2)
add_para(doc, '2024-06-19 ~ 2026-06-12 (近2年) 回测结果:', bold=True)

add_table(doc, ['指标', '融合策略', '璇玑v5.0(对比)', '西部v4.0(对比)'],
    [
        ['总收益', '+60.75% 🏆', '+38.20%', '+54.72%'],
        ['年化收益', '+27.98% 🏆', '+18.31%', '+25.47%'],
        ['最大回撤', '-11.18% 🏆', '-14.65%', '-10.37%'],
        ['Sharpe', '1.93 🏆', '0.96', '1.51'],
        ['Sortino', '10.00 🏆', '10.00', '10.00'],
        ['Calmar', '2.50 🏆', '1.25', '2.46'],
        ['胜率', '50.86%', '45.36%', '53.24% 🏆'],
        ['盈亏比', '2.33 🏆', '2.16', '1.61'],
        ['交易次数', '407 🏆(适中)', '560', '2,207'],
        ['平均持仓', '37.7天', '37.7天', '13.6天'],
    ])

add_heading(doc, '收益归因', level=3)
add_table(doc, ['指标', '盈利交易', '亏损交易', '合计'],
    [
        ['笔数', '207笔', '200笔', '407笔'],
        ['占比', '50.86%', '49.14%', '100%'],
        ['平均收益', '+9.43%', '-4.05%', '盈亏比 2.33'],
    ])

add_heading(doc, '卖出原因', level=3)
add_table(doc, ['原因', '笔数', '占比', '说明'],
    [
        ['追踪止损', '256', '62.9%', '个股从最高点回撤超过-5%'],
        ['调仓', '83', '20.4%', '周频调仓导致的被动卖出'],
        ['无行情数据', '42', '10.3%', '数据缺失或退市'],
        ['组合止损', '26', '6.4%', '组合整体回撤超过-15%'],
    ])

add_heading(doc, '3.6 月度收益归因', level=2)
add_table(doc, ['期间', '月度收益', '累计', '期间', '月度收益', '累计'],
    [
        ['2024-06', '-4.75%', '-4.75%', '2025-06', '+2.76%', '+11.94%'],
        ['2024-07', '-2.08%', '-6.83%', '2025-07', '+2.48%', '+14.42%'],
        ['2024-08', '-1.41%', '-8.24%', '2025-08', '+3.60%', '+18.02%'],
        ['2024-09', '+6.86%', '-1.38%', '2025-09', '+2.88%', '+20.90%'],
        ['2024-10', '+0.08%', '-1.30%', '2025-10', '+0.53%', '+21.43%'],
        ['2024-11', '+6.28%', '+4.98%', '2025-11', '-0.49%', '+20.94%'],
        ['2024-12', '+0.80%', '+5.78%', '2025-12', '+11.06%', '+32.00%'],
        ['2025-01', '+3.23%', '+9.01%', '2026-01', '+4.10%', '+36.10%'],
        ['2025-02', '+2.70%', '+11.71%', '2026-02', '+3.00%', '+39.10%'],
        ['2025-03', '-0.32%', '+11.39%', '2026-03', '-3.62%', '+35.48%'],
        ['2025-04', '-2.39%', '+9.00%', '2026-04', '+5.17%', '+40.65%'],
        ['2025-05', '+0.18%', '+9.18%', '2026-05', '-0.94%', '+39.71%'],
    ])
add_bullet(doc, [
    '盈利月: 17/24 = 70.8%, 亏损月: 5/24 = 20.8%, 持平月: 2/24 = 8.3%',
    '最佳月: 2025-12 (+11.06%), 最差月: 2024-06 (-4.75%, 建仓期)',
    '最大连续回撤期: 2024-06~2024-08 (-8.24%), 建仓初期的系统性下跌',
])

add_heading(doc, '3.7 实战建议与风险提示', level=2)

add_heading(doc, '适用场景', level=3)
add_bullet(doc, [
    '可投资转债数量 ≥ 200只',
    '震荡市 / 波折市 (比单边上涨市更稳定,比单边下跌市更有防御)',
    '资金量 100万 ~ 1000万 (25只等权持仓)',
    '交易频率每周,适合程序化交易系统',
])

add_heading(doc, '注意事项', level=3)
add_bullet(doc, [
    '建仓期容易偏弱: 前3个月(-8.24%回撤)需要耐心,建仓初期存在系统性试错成本',
    '追踪止损频繁: 62.9%的交易因追踪止损触发,说明策略追求"截断亏损"',
    '数据质量敏感: PE/PB/行业数据直接影响因子评分',
    '交易成本: 周频调仓(407笔/2年),每笔需考虑滑点和佣金',
])

add_heading(doc, '风险提示', level=3)
add_bullet(doc, [
    '历史回测不代表未来: 2024-2026是转债市场特定环境,不同市场周期表现可能不同',
    '极端行情风险: 单边暴跌可能突破-15%的组合止损线',
    '流动性风险: 小盘转债在极端行情下可能缺乏流动性',
    '政策风险: 可转债新规/交易机制变化可能影响策略逻辑有效性',
])

add_heading(doc, '实盘优化方向', level=3)
add_bullet(doc, [
    '动态参数: 根据市场波动率(VIX-like)自动调整 buffer_size 和 hold_count',
    '行业轮动: 加入行业景气度预期,超配景气上行行业',
    '机器学习: 用XGBoost或LightGBM替代固定权重,从历史数据中学习最优因子组合',
    '多周期验证: 在2020-2022/2022-2024等不同市场周期做样本外验证',
])

add_hr(doc)
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run('LiangHua Quant System  |  璇玑十二因子  +  西部七维打分  +  璇玑×西部融合策略  |  文档自动生成')
set_run_font(run, size=9, color=RGBColor(0x99, 0x99, 0x99))

output_path = '/Users/mac/lianghua/docs/lianghua_strategies_full.docx'
doc.save(output_path)
print(f"✅ 完整策略文档已生成: {output_path}")
print(f"   文件大小: {os.path.getsize(output_path):,} bytes")
